import requests
import json
import re
import os
from collections import Counter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
from matplotlib import rcParams
from dotenv import load_dotenv
import os
from pathlib import Path

# ===== 字体配置 =====
rcParams['font.sans-serif'] = ['PingFang SC']# 可自行更改为系统已有的字体
rcParams['axes.unicode_minus'] = False

# ===== 配置 =====
URL = "https://api.deepseek.com/v1/chat/completions"
BG = "#0f0f13"
GRAY = "#aaaaaa" 

# ==============================================================
# 基础工具
# ==============================================================
def ask_model(prompt, api_key, url=URL):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}]
    }
    response = requests.post(url, headers=headers, json=data)
    result = response.json()
    return response.json()["choices"][0]["message"]["content"]

def _clean_keywords(keywords: list) -> list:
    stopwords = {
        "非常棒", "非常差", "很好", "很差", "不错", "一般", "好", "差",
        "满意", "不满意", "喜欢", "讨厌", "棒", "赞", "烂", "糟",
        "感觉", "觉得", "认为", "还好", "还行", "可以"
    }
    cleaned, seen = [], set()
    for k in keywords:
        k = k.strip()
        if k and k not in stopwords and len(k) > 1 and k not in seen:
            cleaned.append(k)
            seen.add(k)
    return cleaned[:5]

# ==============================================================
# 数据读取
# ==============================================================
def read_json(file_path: str) -> list:
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    reviews = []
    items = data.get("result", {}).get("items", [])
    for item in items:
        content = item.get("content", "")
        if content and content.strip():
            reviews.append(content.strip())
    return reviews

def read_json_folder(folder_path: str) -> list:
    all_reviews = []
    json_files = sorted([f for f in os.listdir(folder_path) if f.endswith(".json")])
    if not json_files:
        print(f"文件夹 {folder_path} 下未找到 JSON 文件")
        return []
    for fname in json_files:
        path = os.path.join(folder_path, fname)
        reviews = read_json(path)
        print(f"{fname}: 读取 {len(reviews)} 条评论")
        all_reviews.extend(reviews)
    print(f"共读取 {len(all_reviews)} 条原始评论（来自 {len(json_files)} 个文件）")
    return all_reviews

def read_txt(file_path: str) -> list:
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    return [line.strip() for line in lines if line.strip()]

def is_valid_review(text: str) -> bool:
    if len(text) < 5:
        return False
    if len(set(text)) < 3:
        return False
    return True

# ==============================================================
# 模块一：批量真实性过滤
# ==============================================================
def check_authenticity_batch(reviews: list, api_key: str,
                              batch_size: int = 10, threshold: int = 5) -> tuple:
    authentic, suspicious = [], []
    total_batches = (len(reviews) + batch_size - 1) // batch_size

    for i in range(0, len(reviews), batch_size):
        batch = reviews[i:i + batch_size]
        numbered = "\n".join([f"{j+1}. {r}" for j, r in enumerate(batch)])
        prompt = f"""你是反刷评专家。请判断以下 {len(batch)} 条评论是否真实，只输出 JSON 数组：
[
  {{
    "id": 1,
    "authenticity_score": 0到10的整数,
    "is_suspicious": true或false,
    "reason": "一句话"
  }}
]
判断标准：
- 模板化语言（强烈推荐、五星好评等套话）→ 可信度低
- 内容空洞，无具体细节 → 可信度低
- 有具体场景、细节描述 → 可信度高
评论：
{numbered}
只输出 JSON 数组，不要任何解释。"""

        try:
            raw = ask_model(prompt, api_key)
            clean = re.sub(r"```json|```", "", raw).strip()
            items = json.loads(clean)
            for j, item in enumerate(items):
                if j >= len(batch):
                    break
                review = batch[j]
                score = item.get("authenticity_score", 5)
                is_susp = item.get("is_suspicious", False)
                reason = item.get("reason", "")
                if is_susp or score < threshold:
                    suspicious.append((review, {"score": score, "reason": reason}))
                else:
                    authentic.append(review)
        except Exception:
            authentic.extend(batch)
            print(f"真实性过滤第 {i//batch_size + 1} 批解析失败，全部保留")

        print(f"真实性过滤第 {i//batch_size + 1}/{total_batches} 批完成")

    print(f"\n过滤完成：保留 {len(authentic)} 条，剔除 {len(suspicious)} 条可疑评论\n")
    return authentic, suspicious

# ==============================================================
# 模块二：批量情感分析
# ==============================================================
def analyze_sentiment_batch(reviews: list, api_key: str, batch_size: int = 10) -> list:
    all_results = []
    total_batches = (len(reviews) + batch_size - 1) // batch_size

    for i in range(0, len(reviews), batch_size):
        batch = reviews[i:i + batch_size]
        numbered = "\n".join([f"{j+1}. {r}" for j, r in enumerate(batch)])
        prompt = f"""请分析以下 {len(batch)} 条评论的情感，只输出 JSON 数组，格式严格如下：
[
  {{
    "id": 1,
    "sentiment": "positive 或 neutral 或 negative",
    "reason": "一句话说明主要情感原因",
    "keywords": ["关键词1", "关键词2", "关键词3"]
  }}
]
关键词规则：提取 3-5 个描述具体问题或优点的名词/动词短语，禁止纯情感词，禁止语义重复。
评论：
{numbered}
只输出 JSON 数组，不要任何解释或代码块标记。"""

        raw = ask_model(prompt, api_key)
        parsed = _parse_batch_result(raw, batch)

        if len(parsed) != len(batch):
            print(f"第 {i//batch_size + 1} 批返回 {len(parsed)} 条，期望 {len(batch)} 条，逐条重试...")
            parsed = _fallback_one_by_one(batch, api_key)

        all_results.extend(parsed)
        print(f"情感分析第 {i//batch_size + 1}/{total_batches} 批完成（{len(batch)} 条）")

    return all_results

def _parse_batch_result(raw: str, batch: list) -> list:
    try:
        clean = re.sub(r"```json|```", "", raw).strip()
        data = json.loads(clean)
        results = []
        for item in data:
            results.append({
                "sentiment": item.get("sentiment", "unknown"),
                "reason": item.get("reason", ""),
                "keywords": _clean_keywords(item.get("keywords", []))
            })
        return results
    except Exception as e:
        print(f"批量解析失败：{e}，将逐条重试")
        return []

def _fallback_one_by_one(reviews: list, api_key: str) -> list:
    results = []
    for r in reviews:
        prompt = f"""请分析这条评论的情感，只输出 JSON：
{{
  "sentiment": "positive / neutral / negative",
  "reason": "一句话说明主要情感原因",
  "keywords": ["关键词1", "关键词2", "关键词3"]
}}
评论：{r}
只输出 JSON，不要任何解释。"""
        try:
            raw = ask_model(prompt, api_key)
            clean = re.sub(r"```json|```", "", raw).strip()
            item = json.loads(clean)
            results.append({
                "sentiment": item.get("sentiment", "unknown"),
                "reason": item.get("reason", ""),
                "keywords": _clean_keywords(item.get("keywords", []))
            })
        except Exception:
            results.append({"sentiment": "unknown", "reason": "解析失败", "keywords": []})
    return results

# ==============================================================
# 模块三：总结 & 差评归类
# ==============================================================
def summarize_reasons(results: list, api_key: str) -> tuple:
    positive_reviews = []
    negative_reviews = []
    for review, res in results:
        if res["sentiment"] == "positive":
            positive_reviews.append(review)
        elif res["sentiment"] == "negative":
            negative_reviews.append(review)

    def summarize(reviews, label):
        if not reviews:
            return f"没有{label}数据"
        text = "\n".join(reviews[:20])
        prompt = f"""以下是一些用户评论（{label}）：
{text}
请总结最主要的 3-5 个原因，使用简洁中文，每条一行，格式：
- 原因描述"""
        return ask_model(prompt, api_key)

    return summarize(positive_reviews, "好评"), summarize(negative_reviews, "差评")

def classify_negative_reasons(results: list, neg_summary: str, api_key: str) -> dict:
    negative_reviews = [
        review for review, res in results
        if res["sentiment"] == "negative"
    ]
    if not negative_reviews:
        return {}
    text = "\n".join(negative_reviews[:30])
    prompt = f"""已知差评原因总结：
{neg_summary}
差评内容：
{text}
请将每条差评归类到上述原因，统计每个原因对应数量。
只输出 JSON，格式：{{"原因1": 数量, "原因2": 数量}}"""
    result = ask_model(prompt, api_key)
    try:
        clean = re.sub(r"```json|```", "", result).strip()
        start = clean.find("{")
        end = clean.rfind("}") + 1
        return json.loads(clean[start:end])
    except Exception:
        return {}

# ==============================================================
# 模块四：竞品对比
# ==============================================================
def analyze_brand(reviews: list, brand_name: str) -> dict:
    all_keywords = []
    batch_results = analyze_sentiment_batch(reviews, api_key, batch_size=10)
    sentiments = [r["sentiment"] for r in batch_results]
    for r in batch_results:
        all_keywords.extend(r["keywords"])
    counts = Counter(sentiments)
    total = len(sentiments) or 1
    keyword_freq = Counter(all_keywords).most_common(8)
    return {
        "brand": brand_name,
        "total": total,
        "positive_rate": round(counts.get("positive", 0) / total * 100, 1),
        "neutral_rate": round(counts.get("neutral", 0) / total * 100, 1),
        "negative_rate": round(counts.get("negative", 0) / total * 100, 1),
        "top_keywords": keyword_freq,
    }

def compare_with_competitor(my_reviews, competitor_reviews,
                             my_name="我方产品", competitor_name="竞品") -> dict:
    print(f"\n分析 {my_name} 评论...")
    my_profile = analyze_brand(my_reviews, my_name)
    print(f"分析 {competitor_name} 评论...")
    comp_profile = analyze_brand(competitor_reviews, competitor_name)

    prompt = f"""以下是两个产品的用户评论分析数据：
{my_name}：
- 好评率 {my_profile['positive_rate']}%，差评率 {my_profile['negative_rate']}%
- 高频关键词：{[k for k, _ in my_profile['top_keywords']]}
{competitor_name}：
- 好评率 {comp_profile['positive_rate']}%，差评率 {comp_profile['negative_rate']}%
- 高频关键词：{[k for k, _ in comp_profile['top_keywords']]}
请从产品经理视角分析：
1. 我方相比竞品的核心优势是什么？
2. 我方相比竞品的主要差距在哪里？
3. 最值得优先改进的一个方向是什么？
格式：
优势：...
差距：...
优先改进：..."""
    gap_analysis = ask_model(prompt, api_key)
    return {
        "my": my_profile,
        "competitor": comp_profile,
        "gap_analysis": gap_analysis,
    }

# ==============================================================
# 模块五：改进建议
# ==============================================================
def generate_actions(neg_summary: str, reason_counts: dict, competitor_gap: str = None) -> str:
    comp_section = f"\n竞品对比中我方的主要差距：\n{competitor_gap}" if competitor_gap else ""
    sorted_reasons = sorted(reason_counts.items(), key=lambda x: x[1], reverse=True)
    top_reasons = "\n".join([f"- {r}：{c}条" for r, c in sorted_reasons[:5]])
    prompt = f"""你是一名资深产品经理。根据以下用户反馈数据，生成可落地的产品改进建议。
差评核心原因（按数量排序）：
{top_reasons}
差评总结：
{neg_summary}
{comp_section}
请针对排名前 3 的问题，各给出一条具体建议，格式如下：
【问题1】问题描述
→ 建议动作：具体可执行的改进措施（1-2句）
→ 预期效果：改进后用户体验如何变化
【问题2】...
【问题3】...
语言简洁专业，建议要具体，不要泛泛而谈。"""
    return ask_model(prompt, api_key)

# ==============================================================
# 可视化
# ==============================================================
def _style_ax(ax):
    ax.set_facecolor(BG)
    ax.spines[:].set_visible(False)

def _draw_pie(ax, pos, neu, neg):
    total = pos + neu + neg
    pie_data = [x for x in [pos, neu, neg] if x > 0]
    pie_labels = [l for l, x in zip(["正向", "中性", "负向"], [pos, neu, neg]) if x > 0]
    pie_colors = [c for c, x in zip(["#4ade80", "#facc15", "#f87171"], [pos, neu, neg]) if x > 0]
    wedges, texts, autotexts = ax.pie(
        pie_data, labels=pie_labels, colors=pie_colors,
        autopct="%1.1f%%", startangle=140, pctdistance=0.78,
        labeldistance=1.15,
        wedgeprops={"linewidth": 2.5, "edgecolor": BG},
        textprops={"fontsize": 12},
    )
    for t in texts:
        t.set_color("white")
        t.set_fontsize(12)
    for at in autotexts:
        at.set_color(BG)
        at.set_fontsize(10)
        at.set_fontweight("bold")
    ax.text(0, 0, f"{total}\n条评论", ha="center", va="center",
            fontsize=14, color="white", fontweight="bold", linespacing=1.6)
    ax.set_title("情感分布总览", color="white", fontsize=15, pad=20, fontweight="bold")

def _draw_bar(ax, reason_counts):
    if not reason_counts:
        ax.text(0.5, 0.5, "无差评原因数据", transform=ax.transAxes,
                ha="center", va="center", color=GRAY, fontsize=13)
        return
    sorted_items = sorted(reason_counts.items(), key=lambda x: x[1])
    labels_bar = [item[0] for item in sorted_items]
    values_bar = [item[1] for item in sorted_items]
    n = len(labels_bar)
    y_pos = np.arange(n)
    gradient_colors = plt.cm.RdYlGn_r(np.linspace(0.15, 0.75, n))
    bar_height = min(0.55, 0.9 / max(n, 1))
    bars = ax.barh(y_pos, values_bar, color=gradient_colors, height=bar_height, edgecolor="none")
    for bar, val in zip(bars, values_bar):
        ax.text(val + max(values_bar) * 0.02, bar.get_y() + bar.get_height() / 2,
                str(val), va="center", ha="left", color="white", fontsize=10, fontweight="bold")
    wrapped = ["\n".join([lb[i:i+8] for i in range(0, len(lb), 8)]) for lb in labels_bar]
    ax.set_yticks(y_pos)
    ax.set_yticklabels(wrapped, color="white", fontsize=10)
    ax.set_xlabel("评论数量", color=GRAY, fontsize=11, labelpad=8)
    ax.set_title("差评原因分布", color="white", fontsize=15, pad=20, fontweight="bold")
    ax.tick_params(axis="x", colors=GRAY, labelsize=10)
    ax.set_xlim(0, max(values_bar) * 1.35)
    ax.xaxis.grid(True, color="#1e1e2e", linewidth=0.8, linestyle="--")
    ax.set_axisbelow(True)

def visualize(results: list, reason_counts: dict, comparison=None):
    sentiments = [res["sentiment"] for _, res in results]
    counts = Counter(sentiments)
    pos = counts.get("positive", 0)
    neu = counts.get("neutral", 0)
    neg = counts.get("negative", 0)

    # 图1：情感分布饼图
    fig1, ax1 = plt.subplots(figsize=(7, 7))
    fig1.patch.set_facecolor(BG)
    _style_ax(ax1)
    _draw_pie(ax1, pos, neu, neg)
    fig1.tight_layout(pad=2.5)
    fig1.savefig("chart_sentiment_pie.png", dpi=160, bbox_inches="tight", facecolor=BG)
    plt.close(fig1)
    print("情感饼图已保存：chart_sentiment_pie.png")

    # 图2：差评原因条形图
    n_reasons = len(reason_counts)
    fig_h = max(5, n_reasons * 0.7 + 2)
    fig2, ax2 = plt.subplots(figsize=(9, fig_h))
    fig2.patch.set_facecolor(BG)
    _style_ax(ax2)
    _draw_bar(ax2, reason_counts)
    fig2.tight_layout(pad=2.5)
    fig2.savefig("chart_neg_reasons.png", dpi=160, bbox_inches="tight", facecolor=BG)
    plt.close(fig2)
    print("差评原因图已保存：chart_neg_reasons.png")

    # 图3（可选）：竞品对比图
    if comparison:
        my = comparison["my"]
        comp = comparison["competitor"]
        fig3, (ax3, ax4) = plt.subplots(1, 2, figsize=(14, 6))
        fig3.patch.set_facecolor(BG)
        for ax in [ax3, ax4]:
            _style_ax(ax)
        x, w = np.arange(2), 0.32
        b1 = ax3.bar(x - w/2, [my["positive_rate"], comp["positive_rate"]],
                     width=w, label="好评率", color="#4ade80")
        b2 = ax3.bar(x + w/2, [my["negative_rate"], comp["negative_rate"]],
                     width=w, label="差评率", color="#f87171")
        for bar in list(b1) + list(b2):
            h = bar.get_height()
            ax3.text(bar.get_x() + bar.get_width() / 2, h + 1,
                     f"{h}%", ha="center", va="bottom", color="white", fontsize=10, fontweight="bold")
        ax3.set_xticks(x)
        ax3.set_xticklabels([my["brand"], comp["brand"]], color="white", fontsize=11)
        ax3.set_ylabel("百分比 (%)", color=GRAY, fontsize=11)
        ax3.legend(loc="upper right", facecolor=BG, labelcolor="white")
        ax3.set_title("好评率 & 差评率对比", color="white", fontsize=14, pad=15)
        ax3.set_ylim(0, 105)
        ax3.grid(True, color="#1e1e2e", linestyle="--", axis="y")
        my_kws = "、".join([k for k, _ in my["top_keywords"][:5]])
        comp_kws = "、".join([k for k, _ in comp["top_keywords"][:5]])
        ax4.text(0.1, 0.6, f"{my['brand']} 高频关键词\n{my_kws}",
                 transform=ax4.transAxes, color="white", fontsize=12, va="top")
        ax4.text(0.1, 0.2, f"{comp['brand']} 高频关键词\n{comp_kws}",
                 transform=ax4.transAxes, color="white", fontsize=12, va="top")
        ax4.axis("off")
        ax4.set_title("关键词对比", color="white", fontsize=14, pad=15)
        fig3.tight_layout()
        fig3.savefig("chart_competitor.png", dpi=160, bbox_inches="tight", facecolor=BG)
        plt.close(fig3)
        print("竞品对比图已保存：chart_competitor.png")

# ==============================================================
# 写 docx 报告
# ==============================================================
def write_to_docx(results, output_path, pos_summary, neg_summary,
                  action_plan=None, comparison=None, suspicious=None):
    doc = Document()
    title = doc.add_heading("用户评论情感分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、真实性说明
    if suspicious is not None:
        doc.add_heading("一、数据质量说明", level=1)
        doc.add_paragraph(
            f"原始评论经真实性过滤后，共剔除 {len(suspicious)} 条可疑评论，"
            f"保留 {len(results)} 条有效评论用于后续分析。"
        )
        if suspicious:
            doc.add_paragraph("可疑评论示例：")
            for rev, info in suspicious[:3]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[可信度 {info['score']}/10] ").bold = True
                p.add_run(f"{rev[:50]}... → {info['reason']}")

    # 二、情感总结
    doc.add_heading("二、情感分析总结", level=1)
    doc.add_paragraph("【好评原因】").runs[0].bold = True
    doc.add_paragraph(pos_summary)
    doc.add_paragraph("【差评原因】").runs[0].bold = True
    doc.add_paragraph(neg_summary)
    for chart_file, label in [
        ("chart_sentiment_pie.png", "情感分布图"),
        ("chart_neg_reasons.png", "差评原因分布图")
    ]:
        try:
            doc.add_picture(chart_file, width=Inches(5.5))
            doc.add_paragraph("")
        except Exception:
            doc.add_paragraph(f"（{label}未找到，请先运行可视化）")

    # 三、竞品对比
    if comparison:
        doc.add_page_break()
        doc.add_heading("三、竞品对比分析", level=1)
        my = comparison["my"]
        comp = comparison["competitor"]
        doc.add_paragraph(f"{my['brand']}：好评率 {my['positive_rate']}% | 差评率 {my['negative_rate']}%")
        doc.add_paragraph(f"{comp['brand']}：好评率 {comp['positive_rate']}% | 差评率 {comp['negative_rate']}%")
        doc.add_paragraph("差距分析：").runs[0].bold = True
        doc.add_paragraph(comparison["gap_analysis"])

    # 四、改进建议
    if action_plan:
        doc.add_page_break()
        doc.add_heading("四、产品改进建议", level=1)
        doc.add_paragraph(action_plan)

    # 五、逐条明细
    doc.add_page_break()
    doc.add_heading("五、逐条分析明细", level=1)
    sentiment_map = {
        "positive": "✅ 好评", "neutral": "➖ 中性",
        "negative": "❌ 差评", "unknown": "❓ 未知"
    }
    for i, (review, res) in enumerate(results, 1):
        p = doc.add_paragraph()
        p.add_run(f"评论 {i}：").bold = True
        p.add_run(review)
        doc.add_paragraph(f"情感：{sentiment_map.get(res['sentiment'], res['sentiment'])}")
        doc.add_paragraph(f"原因：{res['reason']}")
        doc.add_paragraph(f"关键词：{'、'.join(res['keywords']) if res['keywords'] else '无'}")
        doc.add_paragraph("")

    doc.save(output_path)
    print(f"报告已保存：{output_path}")

# ==============================================================
# 主流程
# ==============================================================
def main():
    #API配置
    load_dotenv(r"C:\Users\Amyxia\Desktop\SP SYSTEM\LLMbased-Sentiment-Report\.env")
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        print("错误：未找到 API Key，请确保 .env 文件中设置了 DEEPSEEK_API_KEY")
        return
    
    # ── 数据源配置 ──
    JSON_FOLDER = r"C:\Users\Amyxia\Desktop\SP SYSTEM\LLMbased-Sentiment-Report\data"      # ← 改成你的 JSON 文件夹路径
    USE_JSON = True              # True=读JSON文件夹，False=读txt
    TXT_FILE = "comments.txt"   # USE_JSON=False 时使用
    

    # 竞品配置（可选，填路径后自动开启）
    competitor_file = None       # 竞品JSON文件夹或txt路径，None=不开启
    competitor_is_json = True
    my_name = "我方产品"
    competitor_name = "竞品A"
    output_file = "result.docx"

    # Step 1: 读取 & 基础过滤
    print("读取评论...")
    reviews = read_json_folder(JSON_FOLDER) if USE_JSON else read_txt(TXT_FILE)
    reviews = [r for r in reviews if is_valid_review(r)]
    print(f"过滤后有效评论：{len(reviews)} 条")

    # Step 2: 真实性过滤
    authentic_reviews, suspicious_reviews = check_authenticity_batch(
        reviews, api_key, batch_size=10, threshold=5
    )

    # Step 3: 情感分析
    print("\n情感分析中...")
    batch_results = analyze_sentiment_batch(authentic_reviews, api_key, batch_size=10)
    results = list(zip(authentic_reviews, batch_results))

    # Step 4: 总结 & 归类
    print("\n生成总结...")
    pos_summary, neg_summary = summarize_reasons(results)
    reason_counts = classify_negative_reasons(results, neg_summary)

    # Step 5: 竞品对比（可选）
    comparison = None
    if competitor_file:
        comp_reviews = read_json_folder(competitor_file) if competitor_is_json else read_txt(competitor_file)
        comp_reviews = [r for r in comp_reviews if is_valid_review(r)]
        comp_authentic, _ = check_authenticity_batch(comp_reviews, api_key, batch_size=10, threshold=5)
        comparison = compare_with_competitor(authentic_reviews, comp_authentic, my_name, competitor_name)
        print("\n竞品差距分析：")
        print(comparison["gap_analysis"])

    # Step 6: 改进建议
    print("\n生成改进建议...")
    competitor_gap = comparison["gap_analysis"] if comparison else None
    action_plan = generate_actions(neg_summary, reason_counts, competitor_gap)
    print("\n改进建议：")
    print(action_plan)

    # Step 7: 可视化 & 报告
    visualize(results, reason_counts, comparison)
    write_to_docx(
        results=results,
        output_path=output_file,
        pos_summary=pos_summary,
        neg_summary=neg_summary,
        action_plan=action_plan,
        comparison=comparison,
        suspicious=suspicious_reviews
    )
    print("\n🎉 全部完成！")

if __name__ == "__main__":
    main()
